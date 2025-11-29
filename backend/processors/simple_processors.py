"""
Simplified processors for demo purposes without external dependencies.
"""
from pathlib import Path
from typing import Dict, Any
import PyPDF2
from PIL import Image
import docx
from backend.utils.logger import processor_logger


class SimplePDFProcessor:
    """Simple PDF processor using only PyPDF2."""

    def __init__(self):
        self.logger = processor_logger

    def process(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF using PyPDF2."""
        self.logger.info(f"Processing PDF: {file_path.name}")

        result = {
            'filename': file_path.name,
            'file_path': str(file_path),
            'file_type': 'pdf',
            'text': '',
            'metadata': {},
            'success': False,
            'error': None
        }

        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text_parts = []

                for page in reader.pages:
                    text_parts.append(page.extract_text())

                result['text'] = '\n'.join(text_parts)
                result['metadata'] = {
                    'page_count': len(reader.pages),
                    'title': reader.metadata.get('/Title', '') if reader.metadata else ''
                }
                result['success'] = True

            self.logger.info(f"Successfully processed PDF: {file_path.name}")

        except Exception as e:
            result['error'] = str(e)
            self.logger.error(f"Error processing PDF {file_path.name}: {e}")

        return result


class SimpleDOCXProcessor:
    """Simple DOCX processor using python-docx."""

    def __init__(self):
        self.logger = processor_logger

    def process(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        self.logger.info(f"Processing DOCX: {file_path.name}")

        result = {
            'filename': file_path.name,
            'file_path': str(file_path),
            'file_type': 'docx',
            'text': '',
            'metadata': {},
            'success': False,
            'error': None
        }

        try:
            doc = docx.Document(str(file_path))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            result['text'] = '\n'.join(text_parts)
            result['metadata'] = {
                'paragraph_count': len(doc.paragraphs),
                'section_count': len(doc.sections)
            }
            result['success'] = True

            self.logger.info(f"Successfully processed DOCX: {file_path.name}")

        except Exception as e:
            result['error'] = str(e)
            self.logger.error(f"Error processing DOCX {file_path.name}: {e}")

        return result


class SimpleImageProcessor:
    """Simple image processor without OCR for demo."""

    def __init__(self):
        self.logger = processor_logger

    def process(self, file_path: Path) -> Dict[str, Any]:
        """Process image file (OCR disabled for demo)."""
        self.logger.info(f"Processing image: {file_path.name}")

        result = {
            'filename': file_path.name,
            'file_path': str(file_path),
            'file_type': 'image',
            'text': 'OCR not available in demo version. Install pytesseract for full functionality.',
            'metadata': {},
            'success': False,
            'error': None
        }

        try:
            image = Image.open(file_path)
            result['metadata'] = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'width': image.width,
                'height': image.height
            }
            result['success'] = True

            self.logger.info(f"Successfully processed image metadata: {file_path.name}")

        except Exception as e:
            result['error'] = str(e)
            self.logger.error(f"Error processing image {file_path.name}: {e}")

        return result


class SimpleAudioProcessor:
    """Simple audio processor without transcription for demo."""

    def __init__(self):
        self.logger = processor_logger

    def process(self, file_path: Path) -> Dict[str, Any]:
        """Process audio file (transcription disabled for demo)."""
        self.logger.info(f"Processing audio: {file_path.name}")

        result = {
            'filename': file_path.name,
            'file_path': str(file_path),
            'file_type': 'audio',
            'text': 'Audio transcription not available in demo version. Install whisper for full functionality.',
            'metadata': {
                'file_extension': file_path.suffix.lower(),
                'file_size': file_path.stat().st_size
            },
            'success': True,
            'error': None
        }

        self.logger.info(f"Successfully processed audio metadata: {file_path.name}")
        return result