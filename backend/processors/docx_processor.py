"""
Microsoft Word document text extraction and processing.
"""
from pathlib import Path
from typing import Dict, Any, List
import docx
from docx.document import Document
from docx.shared import Inches
from backend.utils.logger import processor_logger


class DOCXProcessor:
    """Extract text and metadata from DOCX files."""

    def __init__(self):
        """Initialize DOCX processor."""
        self.logger = processor_logger

    def process(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text and metadata from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary containing extracted text and metadata
        """
        self.logger.info(f"Processing DOCX: {file_path.name}")

        result = {
            'filename': file_path.name,
            'file_path': str(file_path),
            'file_type': 'docx',
            'text': '',
            'paragraphs': [],
            'tables': [],
            'metadata': {},
            'success': False,
            'error': None
        }

        try:
            doc = docx.Document(str(file_path))

            # Extract full text
            full_text = []
            paragraphs = []

            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })

            # Process tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)

                if table_data:
                    tables.append(table_data)
                    # Add table content to full text
                    for row in table_data:
                        full_text.append(' | '.join(row))

            result['text'] = '\n'.join(full_text)
            result['paragraphs'] = paragraphs
            result['tables'] = tables
            result['metadata'] = self._extract_metadata(doc)
            result['success'] = True

            self.logger.info(f"Successfully processed DOCX: {file_path.name}")

        except Exception as e:
            result['error'] = str(e)
            self.logger.error(f"Error processing DOCX {file_path.name}: {e}")

        return result

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract document metadata."""
        metadata = {}

        try:
            core_props = doc.core_properties

            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'comments': core_props.comments or '',
                'last_modified_by': core_props.last_modified_by or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'category': core_props.category or '',
                'language': core_props.language or ''
            }

            # Document statistics
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)

            # Count sections
            metadata['section_count'] = len(doc.sections)

        except Exception as e:
            self.logger.warning(f"Metadata extraction failed: {e}")

        return metadata